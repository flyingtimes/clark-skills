#!/usr/bin/env python3
"""
Generate wireless resource management reports from CSV files (offline mode).
Supported formats: text, excel, html, word.
"""

import argparse
import sys
import os
import pandas as pd
from datetime import datetime

def load_data_from_csv(data_dir='./data'):
    """
    Load wireless resource data from CSV files.
    
    Expected CSV files in data_dir:
    - wr_space_site.csv: Site information
    - wr_sync_rc_enodeb.csv: ENODEB equipment
    - wr_sync_rc_eutrancell.csv: EUTRAN cells
    - wr_sync_rc_aau.csv: AAU equipment
    """
    data = {}
    
    # Define expected files and their display names
    file_map = {
        'wr_space_site.csv': 'sites',
        'wr_sync_rc_enodeb.csv': 'enodeb',
        'wr_sync_rc_eutrancell.csv': 'eutrancell',
        'wr_sync_rc_aau.csv': 'aau',
        'wr_sync_rc_rru.csv': 'rru',
        'wr_sync_rc_ant.csv': 'antenna',
        'wr_sync_rc_bbu.csv': 'bbu'
    }
    
    for filename, key in file_map.items():
        filepath = os.path.join(data_dir, filename)
        if os.path.exists(filepath):
            try:
                data[key] = pd.read_csv(filepath, encoding='utf-8')
                print(f"Loaded {filename}: {len(data[key])} rows")
            except Exception as e:
                print(f"Warning: Could not load {filename}: {e}")
        else:
            print(f"Warning: {filename} not found in {data_dir}")
    
    return data

def generate_summary_from_data(data):
    """Generate summary statistics from loaded data."""
    results = {}
    
    # Site Summary
    if 'sites' in data:
        df = data['sites']
        if 'maintenance_type' in df.columns and 'vip_level' in df.columns:
            site_summary = df.groupby(['maintenance_type', 'vip_level']).size().reset_index(name='site_count')
            site_summary['percentage'] = (site_summary['site_count'] / site_summary['site_count'].sum() * 100).round(2)
            results['Site Summary'] = site_summary
        else:
            results['Site Summary'] = pd.DataFrame({'Note': ['Required columns maintenance_type, vip_level not found']})
    
    # Equipment Summary
    if 'enodeb' in data:
        df = data['enodeb']
        if 'device_type' in df.columns and 'site_id' in df.columns:
            # Filter for active status if column exists
            if 'lifecycle_status' in df.columns:
                df_active = df[df['lifecycle_status'].str.contains('现网有业务', na=False)]
            else:
                df_active = df
            equipment_summary = df_active.groupby('device_type').agg(
                count=('device_type', 'size'),
                sites_covered=('site_id', pd.Series.nunique)
            ).reset_index().sort_values('count', ascending=False)
            results['Equipment Summary'] = equipment_summary
        else:
            results['Equipment Summary'] = pd.DataFrame({'Note': ['Required columns device_type, site_id not found']})
    
    # Cell Technology Distribution
    if 'eutrancell' in data:
        df = data['eutrancell']
        if 'network_technology' in df.columns:
            if 'lifecycle_status' in df.columns:
                df_active = df[df['lifecycle_status'].str.contains('现网有业务', na=False)]
            else:
                df_active = df
            tech_dist = df_active.groupby('network_technology').size().reset_index(name='cell_count').sort_values('cell_count', ascending=False)
            results['Cell Technology Distribution'] = tech_dist
        else:
            results['Cell Technology Distribution'] = pd.DataFrame({'Note': ['Column network_technology not found']})
    
    # Data Quality Status
    quality_issues = []
    
    # Missing coordinates
    if 'sites' in data:
        df = data['sites']
        missing_coords = df[(df['longitude'].isna()) | (df['latitude'].isna())].shape[0]
        quality_issues.append({'issue_type': 'Missing Coordinates', 'issue_count': missing_coords})
    
    # Invalid relationships (cells without enodeb)
    if 'eutrancell' in data and 'enodeb' in data:
        df_cell = data['eutrancell']
        df_enodeb = data['enodeb']
        if 'enodeb_id' in df_cell.columns and 'enodeb_id' in df_enodeb.columns:
            invalid_rels = df_cell[~df_cell['enodeb_id'].isin(df_enodeb['enodeb_id'])].shape[0]
            quality_issues.append({'issue_type': 'Invalid Relationships', 'issue_count': invalid_rels})
    
    # Mandatory field missing (example: AAU installation location)
    if 'aau' in data:
        df = data['aau']
        if 'installation_location' in df.columns:
            if 'lifecycle_status' in df.columns:
                df_active = df[df['lifecycle_status'].str.contains('现网有业务', na=False)]
            else:
                df_active = df
            missing_mandatory = df_active[df_active['installation_location'].isna()].shape[0]
            quality_issues.append({'issue_type': 'Mandatory Field Missing', 'issue_count': missing_mandatory})
    
    if quality_issues:
        results['Data Quality Status'] = pd.DataFrame(quality_issues)
    else:
        results['Data Quality Status'] = pd.DataFrame({'Note': ['No data quality checks performed due to missing data']})
    
    return results

def generate_text_report(results, timestamp, output_file=None):
    """Generate text format report."""
    report_lines = []
    report_lines.append("=" * 60)
    report_lines.append("WIRELESS RESOURCE MANAGEMENT REPORT (OFFLINE)")
    report_lines.append(f"Generated: {timestamp}")
    report_lines.append("=" * 60)
    
    for report_name, df in results.items():
        report_lines.append(f"\n{report_name}")
        report_lines.append("-" * 40)
        
        if df.empty:
            report_lines.append("No data available")
        else:
            df_string = df.to_string(index=False)
            report_lines.append(df_string)
    
    report_lines.append("\n" + "=" * 60)
    report_lines.append("END OF REPORT")
    
    report_content = "\n".join(report_lines)
    
    # Output to file or console
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(report_content)
        print(f"Text report saved to {output_file}")
    else:
        print(report_content)
    
    return True

def generate_excel_report(results, timestamp, output_file=None):
    """Generate Excel format report."""
    if not output_file:
        output_file = f"wireless_report_offline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    elif not output_file.endswith('.xlsx'):
        output_file = output_file + '.xlsx'
    
    try:
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            for report_name, df in results.items():
                # Clean sheet name
                sheet_name = report_name[:31].replace(':', '').replace('\\', '').replace('/', '').replace('?', '').replace('*', '').replace('[', '').replace(']', '')
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"Excel report saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error generating Excel report: {e}")
        return False

def generate_html_report(results, timestamp, output_file=None):
    """Generate HTML format report."""
    if not output_file:
        output_file = f"wireless_report_offline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    elif not output_file.endswith('.html'):
        output_file = output_file + '.html'
    
    try:
        html_parts = []
        html_parts.append("""<!DOCTYPE html>
<html>
<head>
    <title>Wireless Resource Management Report (Offline)</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
        h2 { color: #34495e; margin-top: 30px; }
        table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
        th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        th { background-color: #3498db; color: white; }
        tr:nth-child(even) { background-color: #f2f2f2; }
        .timestamp { color: #7f8c8d; font-style: italic; }
        .footer { margin-top: 40px; color: #95a5a6; text-align: center; }
    </style>
</head>
<body>
""")
        html_parts.append(f'<h1>Wireless Resource Management Report (Offline)</h1>')
        html_parts.append(f'<p class="timestamp">Generated: {timestamp}</p>')
        
        for report_name, df in results.items():
            html_parts.append(f'<h2>{report_name}</h2>')
            if df.empty:
                html_parts.append('<p>No data available</p>')
            else:
                html_table = df.to_html(index=False, classes='data-table')
                html_parts.append(html_table)
        
        html_parts.append("""
    <div class="footer">
        <p>End of report. Generated from offline CSV data.</p>
    </div>
</body>
</html>
""")
        
        html_content = '\n'.join(html_parts)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print(f"HTML report saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error generating HTML report: {e}")
        return False

def generate_word_report(results, timestamp, output_file=None):
    """Generate Word format report."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx package not installed. Install with: pip install python-docx")
        return False
    
    if not output_file:
        output_file = f"wireless_report_offline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    elif not output_file.endswith('.docx'):
        output_file = output_file + '.docx'
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Wireless Resource Management Report (Offline)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        doc.add_paragraph(f'Generated: {timestamp}')
        doc.add_paragraph()
        
        # Add each section
        for report_name, df in results.items():
            doc.add_heading(report_name, level=2)
            
            if df.empty:
                doc.add_paragraph('No data available')
            else:
                table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                table.style = 'LightShading-Accent1'
                
                # Header row
                for col_idx, col_name in enumerate(df.columns):
                    cell = table.cell(0, col_idx)
                    cell.text = str(col_name)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Data rows
                for row_idx, row in df.iterrows():
                    for col_idx, col_name in enumerate(df.columns):
                        table.cell(row_idx+1, col_idx).text = str(row[col_name])
            
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('End of report. Generated from offline CSV data.')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(output_file)
        print(f"Word report saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error generating Word report: {e}")
        return False

def main():
    parser = argparse.ArgumentParser(description='Generate wireless resource management reports from CSV files (offline)')
    parser.add_argument('--data-dir', '-d', default='./data', help='Directory containing CSV files (default: ./data)')
    parser.add_argument('--output', '-o', help='Output file path (default: based on format)')
    parser.add_argument('--format', '-f', choices=['text', 'excel', 'html', 'word'], default='text',
                       help='Output format (text, excel, html, or word)')
    
    args = parser.parse_args()
    
    # Check if data directory exists
    if not os.path.exists(args.data_dir):
        print(f"Error: Data directory '{args.data_dir}' not found.")
        print("Please create the directory and place CSV files there, or specify with --data-dir.")
        sys.exit(1)
    
    # Load data from CSV
    print(f"Loading data from {args.data_dir}...")
    data = load_data_from_csv(args.data_dir)
    
    if not data:
        print("Error: No CSV files could be loaded. Ensure CSV files exist in the data directory.")
        sys.exit(1)
    
    # Generate summary statistics
    results = generate_summary_from_data(data)
    
    if not results:
        print("Error: Could not generate any summary statistics.")
        sys.exit(1)
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Generate report in specified format
    if args.format == 'text':
        success = generate_text_report(results, timestamp, args.output)
    elif args.format == 'excel':
        success = generate_excel_report(results, timestamp, args.output)
    elif args.format == 'html':
        success = generate_html_report(results, timestamp, args.output)
    elif args.format == 'word':
        success = generate_word_report(results, timestamp, args.output)
    else:
        print(f"Error: Unsupported format '{args.format}'")
        sys.exit(1)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()