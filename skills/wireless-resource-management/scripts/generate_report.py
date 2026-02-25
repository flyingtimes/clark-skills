#!/usr/bin/env python3
"""
Generate wireless resource management reports in multiple formats.
Supported formats: text, excel, html, word.
"""

import argparse
import sys
import os
from datetime import datetime
from db_config import config

def generate_resource_summary(output_file=None, format='text'):
    """Generate resource summary report in specified format."""
    
    try:
        import psycopg2
        import pandas as pd
        conn = psycopg2.connect(**config.psycopg2_params())
        
        # Collect summary statistics
        queries = {
            "Site Summary": """
                SELECT 
                  maintenance_type,
                  vip_level,
                  COUNT(*) as site_count,
                  ROUND(100.0 * COUNT(*) / SUM(COUNT(*)) OVER (), 2) as percentage
                FROM wr_space_site
                WHERE maintenance_type IS NOT NULL
                GROUP BY maintenance_type, vip_level
                ORDER BY site_count DESC
            """,
            "Equipment Summary": """
                SELECT 
                  device_type,
                  COUNT(*) as count,
                  COUNT(DISTINCT site_id) as sites_covered
                FROM wr_sync_rc_enodeb
                WHERE lifecycle_status = '现网有业务'
                GROUP BY device_type
                ORDER BY count DESC
            """,
            "Cell Technology Distribution": """
                SELECT 
                  network_technology,
                  COUNT(*) as cell_count
                FROM wr_sync_rc_eutrancell
                WHERE lifecycle_status = '现网有业务'
                GROUP BY network_technology
                ORDER BY cell_count DESC
            """,
            "Data Quality Status": """
                SELECT 
                  'Missing Coordinates' as issue_type,
                  COUNT(*) as issue_count
                FROM wr_space_site
                WHERE longitude IS NULL OR latitude IS NULL
                UNION ALL
                SELECT 
                  'Invalid Relationships',
                  COUNT(*)
                FROM wr_sync_rc_eutrancell c
                LEFT JOIN wr_sync_rc_enodeb e ON c.enodeb_id = e.enodeb_id
                WHERE e.enodeb_id IS NULL
                UNION ALL
                SELECT 
                  'Mandatory Field Missing',
                  COUNT(*)
                FROM wr_sync_rc_aau
                WHERE installation_location IS NULL 
                  AND lifecycle_status = '现网有业务'
            """
        }
        
        # Execute queries and collect results
        results = {}
        for report_name, query in queries.items():
            df = pd.read_sql_query(query, conn)
            results[report_name] = df
        
        conn.close()
        
        # Generate report based on format
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        if format == 'text':
            return generate_text_report(results, timestamp, output_file)
        elif format == 'excel':
            return generate_excel_report(results, timestamp, output_file)
        elif format == 'html':
            return generate_html_report(results, timestamp, output_file)
        elif format == 'word':
            return generate_word_report(results, timestamp, output_file)
        else:
            print(f"Error: Unsupported format '{format}'")
            return False
        
    except ImportError as e:
        print(f"Error: Required package not installed: {e}")
        print("Install with: pip install pandas psycopg2-binary openpyxl")
        if format == 'word':
            print("For Word output, also install: pip install python-docx")
        return False
    except Exception as e:
        print(f"Error generating report: {e}")
        return False

def generate_text_report(results, timestamp, output_file=None):
    """Generate text format report."""
    report_lines = []
    report_lines.append("=" * 60)
    report_lines.append("WIRELESS RESOURCE MANAGEMENT REPORT")
    report_lines.append(f"Generated: {timestamp}")
    report_lines.append("=" * 60)
    
    for report_name, df in results.items():
        report_lines.append(f"\n{report_name}")
        report_lines.append("-" * 40)
        
        if df.empty:
            report_lines.append("No data available")
        else:
            df_string = df.to_string(index=False)
            report_lines.append(df_string)
    
    report_lines.append("\n" + "=" * 60)
    report_lines.append("END OF REPORT")
    
    report_content = "\n".join(report_lines)
    
    # Output to file or console
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(report_content)
        print(f"Text report saved to {output_file}")
    else:
        print(report_content)
    
    return True

def generate_excel_report(results, timestamp, output_file=None):
    """Generate Excel format report."""
    import pandas as pd
    
    if not output_file:
        output_file = f"wireless_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    elif not output_file.endswith('.xlsx'):
        output_file = output_file + '.xlsx'
    
    try:
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            for report_name, df in results.items():
                # Clean sheet name (Excel limit 31 chars, no special chars)
                sheet_name = report_name[:31].replace(':', '').replace('\\', '').replace('/', '').replace('?', '').replace('*', '').replace('[', '').replace(']', '')
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"Excel report saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error generating Excel report: {e}")
        return False

def generate_html_report(results, timestamp, output_file=None):
    """Generate HTML format report."""
    import pandas as pd
    
    if not output_file:
        output_file = f"wireless_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    elif not output_file.endswith('.html'):
        output_file = output_file + '.html'
    
    try:
        # Create HTML content
        html_parts = []
        html_parts.append("""<!DOCTYPE html>
<html>
<head>
    <title>Wireless Resource Management Report</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
        h2 { color: #34495e; margin-top: 30px; }
        table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
        th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        th { background-color: #3498db; color: white; }
        tr:nth-child(even) { background-color: #f2f2f2; }
        .timestamp { color: #7f8c8d; font-style: italic; }
        .footer { margin-top: 40px; color: #95a5a6; text-align: center; }
    </style>
</head>
<body>
""")
        html_parts.append(f'<h1>Wireless Resource Management Report</h1>')
        html_parts.append(f'<p class="timestamp">Generated: {timestamp}</p>')
        
        for report_name, df in results.items():
            html_parts.append(f'<h2>{report_name}</h2>')
            if df.empty:
                html_parts.append('<p>No data available</p>')
            else:
                html_table = df.to_html(index=False, classes='data-table')
                html_parts.append(html_table)
        
        html_parts.append("""
    <div class="footer">
        <p>End of report. This report contains confidential wireless resource data.</p>
    </div>
</body>
</html>
""")
        
        html_content = '\n'.join(html_parts)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print(f"HTML report saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error generating HTML report: {e}")
        return False

def generate_word_report(results, timestamp, output_file=None):
    """Generate Word format report."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx package not installed. Install with: pip install python-docx")
        return False
    
    if not output_file:
        output_file = f"wireless_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    elif not output_file.endswith('.docx'):
        output_file = output_file + '.docx'
    
    try:
        # Create a new Document
        doc = Document()
        
        # Title
        title = doc.add_heading('Wireless Resource Management Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        doc.add_paragraph(f'Generated: {timestamp}')
        doc.add_paragraph()
        
        # Add each section
        for report_name, df in results.items():
            # Add section heading
            doc.add_heading(report_name, level=2)
            
            if df.empty:
                doc.add_paragraph('No data available')
            else:
                # Create table
                table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                table.style = 'LightShading-Accent1'
                
                # Header row
                for col_idx, col_name in enumerate(df.columns):
                    cell = table.cell(0, col_idx)
                    cell.text = str(col_name)
                    # Make header bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Data rows
                for row_idx, row in df.iterrows():
                    for col_idx, col_name in enumerate(df.columns):
                        table.cell(row_idx+1, col_idx).text = str(row[col_name])
            
            doc.add_paragraph()  # Add space between sections
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph('End of report. This report contains confidential wireless resource data.')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        doc.save(output_file)
        print(f"Word report saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error generating Word report: {e}")
        return False

def main():
    parser = argparse.ArgumentParser(description='Generate wireless resource management reports')
    parser.add_argument('--output', '-o', help='Output file path (default: based on format)')
    parser.add_argument('--format', '-f', choices=['text', 'excel', 'html', 'word'], default='text',
                       help='Output format (text, excel, html, or word)')
    
    args = parser.parse_args()
    
    success = generate_resource_summary(args.output, args.format)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()